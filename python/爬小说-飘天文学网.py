#! python3
#!/usr/bin/env python
# -*- coding: utf-8 -*-
# @Version      : 2021.10.13.2
# @Author       : YIU
# @Github       : https://github.com/usaginya
# @For website  : www.piaotianwenxue.com


import ctypes
import os
import re
import random
import sys
import subprocess


def checkPackages(usePackages):
    '''包检查

    检查需要的包是否已安装

    Args:
        usePackages (string[]): 需要的第三方包名数组

    Returns:
        bool: 全部已安装返回 True 否则为 False
    '''
    if not usePackages or len(usePackages) < 1:
        return False

    print('\n' + '='*78)
    print('\n\n\t正在检查需要的包、请稍等...\n\n')

    reqs = subprocess.check_output([sys.executable, '-m', 'pip', 'freeze'])
    installed_packages = [r.decode().split('==')[0] for r in reqs.split()]

    demand = ''
    for packName in usePackages:
        if not packName in installed_packages:
            demand = f'{demand}pip3 install {packName}\n'
    if demand:
        demand = f'pip3 install --upgrade pip --user\n{demand}'
        print('\n请先在 cmd 粘贴以下命令安装需要的包、安装完成后再使用本脚本\n\n')
        print(demand)
        print('='*78 + '\n\n')
        return False

    print('\n\t\t√ 没有问题\n\n' + '='*78 + '\n\n')
    return True


# 检查需要的包
if not checkPackages(['beautifulsoup4', 'python-docx', 'requests', 'urllib3']):
    input('按任意键退出脚本...')
    sys.exit()


# ~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~

from bs4 import BeautifulSoup
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt, RGBColor
from docx.oxml.ns import qn
from docx import Document
import requests
import urllib3

USER_AGENT_LIST = [
    "Mozilla/5.0 (Windows NT 6.1; WOW64) AppleWebKit/537.1 (KHTML, like Gecko) Chrome/22.0.1207.1 Safari/537.1",
    "Mozilla/5.0 (X11; CrOS i686 2268.111.0) AppleWebKit/536.11 (KHTML, like Gecko) Chrome/20.0.1132.57 Safari/536.11",
    "Mozilla/5.0 (Windows NT 6.1; WOW64) AppleWebKit/536.6 (KHTML, like Gecko) Chrome/20.0.1092.0 Safari/536.6",
    "Mozilla/5.0 (Windows NT 6.2) AppleWebKit/536.6 (KHTML, like Gecko) Chrome/20.0.1090.0 Safari/536.6",
    "Mozilla/4.0 (compatible; MSIE 7.0; Windows NT 5.1; 360SE)",
    "Mozilla/5.0 (Windows NT 6.1; WOW64) AppleWebKit/536.3 (KHTML, like Gecko) Chrome/19.0.1061.1 Safari/536.3",
    "Mozilla/5.0 (Windows NT 6.1) AppleWebKit/536.3 (KHTML, like Gecko) Chrome/19.0.1061.1 Safari/536.3",
    "Mozilla/5.0 (Windows NT 6.2) AppleWebKit/536.3 (KHTML, like Gecko) Chrome/19.0.1061.0 Safari/536.3",
    "Mozilla/5.0 (X11; Linux x86_64) AppleWebKit/535.24 (KHTML, like Gecko) Chrome/19.0.1055.1 Safari/535.24",
    "Mozilla/5.0 (Windows NT 6.2; WOW64) AppleWebKit/535.24 (KHTML, like Gecko) Chrome/19.0.1055.1 Safari/535.24"
]


class ColorPrint:
    __FOREGROUND_DARKWHITE = 0x07
    __FOREGROUND_GREEN = 0x0a
    __FOREGROUND_SKYBLUE = 0x0b
    __FOREGROUND_PINK = 0x0d
    __FOREGROUND_YELLOW = 0x0e
    __std_out_handle = 0

    def __init__(self):
        self.__std_out_handle = ctypes.windll.kernel32.GetStdHandle(-11)

    def set_cmd_text_color(self, color):
        Bool = ctypes.windll.kernel32.SetConsoleTextAttribute(
            self.__std_out_handle, color)
        return Bool

    def resetColor(self):
        self.set_cmd_text_color(self.__FOREGROUND_DARKWHITE)

    def __printEnd(self, msg):
        sys.stdout.write(msg)
        self.resetColor()

    def yellow(self, msg):
        self.set_cmd_text_color(self.__FOREGROUND_YELLOW)
        self.__printEnd(msg)

    def green(self, msg):
        self.set_cmd_text_color(self.__FOREGROUND_GREEN)
        self.__printEnd(msg)

    def skyblue(self, msg):
        self.set_cmd_text_color(self.__FOREGROUND_SKYBLUE)
        self.__printEnd(msg)

    def pink(self, msg):
        self.set_cmd_text_color(self.__FOREGROUND_PINK)
        self.__printEnd(msg)
# ColorPrint


class Options:
    def __init__(self, bookUrl=None, merge=0, useFormat=True, saveType=0):
        '''输入的参数配置

            Args:
                bookUrl (str):      小说目录链接地址
                merge (int):        章节合并数
                useFormat (bool):   使用排版
                saveType (int):     保存格式 [0:txt, 1:docx]
        '''
        self.bookUrl = bookUrl
        self.merge = merge if merge < 10000 else 10000
        self.useFormat = useFormat
        self.saveType = saveType
# Options


# 初始化全局变量
options = Options()
cprint = ColorPrint()
s = requests.Session()
urllib3.disable_warnings(urllib3.exceptions.InsecureRequestWarning)


def toInt(str):
    '''字符串转为整数

    Returns:
        int: 返回对应整数类型、失败时返回 0
    '''
    try:
        int(str)
        return int(str)
    except ValueError:  # 报类型错误、说明不是int字符串
        try:
            float(str)  # 验证是不是浮点字符串
            return int(float(str))
        except ValueError:  # 即不是浮点、也不是int字符串
            return 0


def zeroInt(intNumber, zeroNumber=5):
    '''整数补零转为字符串

    Args:
        intNumber (int): 需要补零的整数
        zeroNumber (int): 补零位数

    Returns:
        str: 返回补零后的字符串
    '''
    return str(intNumber).zfill(zeroNumber)


def showScriptInfo():
    '''显示脚本信息'''
    os.system('cls')
    cprint.skyblue('''
    ╔╦╗ ┌─┐ ┌─┐ ┬┌─ ┌─┐ ┬   ╔═╗ ┬ ┬ ┌┬┐ ┬ ┬ ┌─┐ ┌┐┌
    ║║║ │ │ ├┤  ├┴┐ ├─┤ │   ╠═╝ └┬┘  │  ├─┤ │ │ │││
    ╩ ╩ └─┘ └─┘ ┴ ┴ ┴ ┴ ┴   ╩    ┴   ┴  ┴ ┴ └─┘ ┘└┘
    ''')


def inputUrl():
    '''要求输入小说目录页面链接地址

    Returns:
        String: 小说目录页面链接地址
    '''
    cprint.skyblue('\n请粘贴小说的目录页面链接地址，然后回车：')
    bookUrl = input()
    bookUrl = re.search(
        r'^(http|https):\/\/www.piaotianwenxue.com/book\/\d+\/\d+\/?$', bookUrl, re.M | re.I)

    if not bookUrl:
        showScriptInfo()
        cprint.yellow(
            '\n链接地址不正确，必须是飘天文学网(www.piaotianwenxue.com)的小说目录页面\n\n\n< 重试 >')
        bookUrl = inputUrl()
        if isinstance(bookUrl, str):
            return bookUrl
    return bookUrl.group()


def inputMerge():
    '''要求输入章节合并数

    每 n 章节合并为一个文件

    Returns:
        int: 章节合并数 无合并数返回 0
    '''
    cprint.skyblue('\n请输入把几章合并为一章、不输入或者输入的不是数字则每章单独保存：')
    result = input()
    result = re.search(r'^\d+$', result, re.M | re.I)
    if result:
        cprint.green(f'\n [ 把 {result.group()} 章合并为一章 ]\n')
        return toInt(result.group())
    cprint.green(f'\n [ 每章单独保存 ]\n')
    return 0


def inputFormat():
    '''要求输入确认是否启用排版

    Returns:
        Bool: 是否启用排版
    '''
    cprint.skyblue('\n是否禁用内容排版？当前启用、输入任意内容回车禁用：')
    result = input()
    if result:
        cprint.green('\n [ 排版已禁用 ]\n')
        return False
    cprint.green('\n [ 排版已启用 ]\n')
    return True


def inputSaveType():
    '''要求输入保存格式

    Returns:
        int: 0:txt, 1:docx
    '''
    cprint.skyblue('\n选择要保存的格式、默认保存为 txt 格式、输入任何内容回车则保存为 docx 文档格式：')
    result = input()
    if result:
        cprint.green(f'\n [ 保存为 docx 文档格式 ]\n')
        return 1
    cprint.green(f'\n [ 保存为 txt 格式 ]\n')
    return 0


def inputRequest():
    '''请求输入参数信息

    按步骤将需要输入的信息输入到 options 设置中、并显示最终输入的结果
    '''
    # 顺序请求输入
    options.bookUrl = inputUrl()
    options.merge = inputMerge()
    options.useFormat = inputFormat()
    options.saveType = inputSaveType()
    # 显示输入信息结果
    showScriptInfo()
    cprint.green('\n' + '='*78 + '\n\n')

    cprint.yellow(f'\t小说目录链接： {options.bookUrl}\n')

    cprint.yellow('\t章节合并处理： ')
    if options.merge > 0:
        cprint.yellow(f'把 {options.merge} 章合并为一章\n')
    else:
        cprint.yellow('每章单独保存\n')

    tmpTip = '√ 使用' if options.useFormat else 'X 不用'
    cprint.yellow(f'\t使用排版格式： {tmpTip}\n')

    tmpTip = 'docx 文档' if options.saveType > 0 else 'txt 文本文件'
    cprint.yellow(f'\t保存小说格式： {tmpTip}\n')

    cprint.green('\n' + '='*78 + '\n\n')


def getHtml(url):
    '''请求获取 HTML 内容

    给个链接返回 HTML 内容字符串

    Args:
        url (str): 网页链接

    Returns:
        String: HTML 内容
    '''
    headers = {
        'Accept': 'text/html,application/xhtml+xml,application/xml;q=0.9',
        'Accept-Encoding': 'gzip, deflate',
        'Accept-Language': 'zh-CN,zh;q=0.8',
        'Host': re.search("https?://(?P<host>[^\/\s]+)", url).group("host"),
        'User-Agent': random.choice(USER_AGENT_LIST)
    }

    try:
        response = s.get(url=url, headers=headers, verify=False)
    except:
        return ''

    if response.status_code == 200:
        return response.text
    return ''


def writeTxt(folderName, fileName, content):
    '''写文本到指定文件夹内的 txt 文件

    把小说文本保存到指定文件夹内的 txt 文件里

    Args:
        folderName (str): 文件夹名称、小说名称
        fileName (str): 文件名、章节名
        content (str): 文本内容、小说内容
        skipBreak (bool): 跳过添加分页、用于最后一章、避免有空白页
    '''
    if not os.path.exists(folderName):
        os.makedirs(folderName)

    with open(f'{folderName}\\{fileName}.txt', 'w', encoding='utf-8') as f:
        f.write(content)


def writeDocx(document, title, content, skipBreak=False):
    '''写文本到 Document 对象

    把文本内容写到 Document 对象、写好最后再用 saveDocx 保存

    Args:
        document: Document 对象
        title (str): 标题
        content (str): 文本内容、小说内容
    '''
    fontName = 'SimSun'
    document.styles['Normal'].font.name = fontName
    document.styles['Normal']._element.rPr.rFonts.set(
        qn('w:eastAsia'), fontName)
    document.styles['Normal'].font.size = Pt(12)

    # 写标题
    title_ = document.add_heading()

    if options.useFormat:
        title_.alignment = WD_ALIGN_PARAGRAPH.CENTER

    title_run = title_.add_run(title)

    if options.useFormat:
        title_run.font.size = Pt(18)

    title_run.font.name = fontName
    title_run.element.rPr.rFonts.set(qn('w:eastAsia'), fontName)
    title_run.font.color.rgb = RGBColor(0, 0, 0)

    # 写内容
    document.add_paragraph(content)
    if options.useFormat and not skipBreak:
        document.add_page_break()


def saveDocx(document, folderName, fileName):
    '''保存 Document 到指定文件夹内的 docx 文件

    Args:
        document: Document 对象
        folderName (str): 文件夹名称、小说名称
        fileName (str): 文件名、章节名
    '''
    if not os.path.exists(folderName):
        os.makedirs(folderName)
    try:
        document.save(f'{folderName}\\{fileName}.docx')
    except:
        cprint.pink(f'\n{folderName}\\{fileName}.docx 文件可能被占用、保存失败...')


def getNovelInfo(noveUrl):
    '''获取小说信息

    获取小说信息并返回一个字典

    Args:
        noveUrl (str): 小说详情页面链接地址

    Returns:
        Dictionary: 返回小说信息字典 失败返回 None
            title (str): 小说标题
            catalogs (array): 目录信息组
                .string: 得到目录标题
                ['href']: 得到目录内容链接
    '''
    data = getHtml(noveUrl)
    if not data.strip():
        return None

    soup = BeautifulSoup(data, 'html.parser')

    # 小说标题
    title = soup.select_one('.bookTitle')
    title = title.string if title else ''

    # 小说章节目录信息
    catalogs = soup.select('.panel-chapterlist:last-child a')
    if not catalogs:
        return None

    # 小说信息字典：标题、章节信息组[i]目录标题(.string), 目录链接(['href'])
    dictNovel = {'title': title, 'catalogs': catalogs}
    return dictNovel


def getNovelContent(noveUrl, noveTitle=None):
    '''获取小说内容

    通过小说章节标题和小说内容页链接地址、获取含章节标题的小说内容

    Args:
        noveUrl (str): 小说内容页面链接地址
        noveTitle (str): 小说章节标题、默认只返回小说内容

    Returns:
        String: 返回含章节标题的小说内容 失败返回 None
    '''
    data = getHtml(noveUrl)
    if not data.strip():
        return None
    # 解析html
    soup = BeautifulSoup(getHtml(noveUrl), 'html.parser')
    content = soup.select_one('#htmlContent')
    if content:
        lines = content.get_text().splitlines()
        # 内容排版 段落加空格
        if options.useFormat:
            content = '\n    '.join(lines)
        else:
            content = '\n'.join(lines)
        # 把标题添加到内容前
        return content if not noveTitle else f'{noveTitle}\n{content}'
    return None


def climb():
    '''给我爬

    爬爬爬  我最会爬了

    Returns:
        String: 返回小说标题
    '''
    # 请求输入信息
    inputRequest()
    # 确认信息
    cprint.pink('请确认以上设置是否符合你的需求、如果不满足需求、请重新运行脚本\n')
    cprint.pink('\n\t\t按任意键开始爬取小说...\n')
    input()

    # 简单验证信息防止强行结束错误
    if not options.bookUrl:
        cprint.pink('\n\t哦豁...链接不对')
        return False

    # 开始爬取
    showScriptInfo()
    cprint.skyblue('\n\n' + '='*78 + '\n\n')
    cprint.yellow('\t\t\t>ω0  爬取开始！')

    # 获取小说标题、章节列表信息
    novel = getNovelInfo(options.bookUrl)
    if not novel:
        cprint.pink('\n\t小说信息获取失败...没救了')
        return False

    bookTitle = novel['title']
    catalogs = novel['catalogs']
    mergeCount = 0      # 合并章节计数
    getCount = 0        # 已获得章节计数
    groupCount = 0      # 章节分组计数
    startChapter = 1    # 上次保存章节数记录
    txtContent = ''     # 保存txt的文章内容
    document = Document()   # 创建初始 docx 文档

    # 获取小说章节信息并爬取内容保存到本地文件
    for catalog in catalogs:
        # 获取章节信息：章节标题、章节链接
        title = catalog.string
        novelUrl = catalog['href']

        # 显示爬取信息
        cprint.green(f'\n\n\t正在爬《{bookTitle}》章节：{title} ...')

        if not novelUrl:
            errInfo = f'{title} 章节信息获取失败...' if title else '当前章节信息获取失败...'
            cprint.yellow(f'\n\tX {errInfo}')
            continue

        # 提取章节链接并抓章节内容
        novelHost = re.search(
            '(?P<mainHost>https?://[^\/\s]+)', options.bookUrl).group('mainHost')
        novelUrl = f'{novelHost}{novelUrl}' if novelUrl[0] == '/' else novelUrl

        # 爬取内容
        content = ''
        if options.saveType > 0:
            content = getNovelContent(novelUrl)
        else:
            content = getNovelContent(novelUrl, title)

        # 检查内容
        if not content:
            cprint.yellow('\n\tX 章节内容获取失败...')
            continue

        # 合并章节数+1
        mergeCount += 1
        # 已获得章节数+1
        getCount += 1

        # 写内容
        if options.saveType > 0:  # docx
            writeDocx(document, title=title, content=content,
                      skipBreak=mergeCount >= options.merge)
        else:  # txt
            txtContent += ('\n'*5 if mergeCount > 1 else txtContent) + content

        # 保存文件
        if mergeCount >= options.merge:
            # 章节分组+1
            groupCount += 1

            # 格式化文件名
            fname = ''
            if options.merge > 0:  # 使用了章节合并
                # 0001 小说名-章节00001-00100
                fname = f'{zeroInt(groupCount,4)} {bookTitle}-章节{zeroInt(startChapter)}-{zeroInt(getCount)}'

            else:  # 不使用章节合并
                fname = re.sub(r'[\\/:*?"<>|\r\n]+', '', title)
                fname = zeroInt(groupCount) if not fname.strip() else fname

            # 根据保存格式保存文件
            if options.saveType > 0:  # docx
                saveDocx(document, folderName=bookTitle, fileName=fname)
                document = Document()  # 创建新文档

            else:  # txt
                writeTxt(folderName=bookTitle,
                         fileName=fname, content=txtContent)

            mergeCount = 0
            txtContent = ''
            startChapter = getCount + 1    # 记录章节位置
            cprint.skyblue('\n\n\t√ 以上章节保存成功')
        # ------
    # --- for end
    return bookTitle


# --- 开始主线任务 ---
# 显示信息
showScriptInfo()
cprint.skyblue('\n\n' + '='*78 + '\n\n')

# 任务结束
novelTitle = climb()
showScriptInfo()
cprint.green('\n' + '='*78 + '\n')
cprint.skyblue(f'\n\t《{novelTitle}》 小说爬取完成！\n')
cprint.skyblue(f'\n\t按任意键退出脚本...\n')
cprint.green('\n' + '='*78 + '\n')
input()
sys.exit()
